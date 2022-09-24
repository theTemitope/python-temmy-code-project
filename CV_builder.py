#building a CV
from docx import Document
from  docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('Mine.jpg', width=Inches(1.0))

# name, phone number and email address
name = input('What is your name? ')
speak('Hello ' + name + 'How are you today?')
speak('What is your Phone_number?')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

p = document.add_paragraph()
p.add_run(name + ' | ' + phone_number + ' | ' + email).bold = True

# about me
document.add_heading('About me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)


# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()
company = input('Enter Company: ')
start_date = input('Start Date: ')
to_date = input('To Date: ')
experience = input('Describe your experience at '+ company + ': ')

p.add_run(company + ' ' + '\n').bold = True
p.add_run(start_date + '-' + to_date + '\n').italics = True
p.add_run(experience)

# more experiences
while True:
    more_experiences = input(
        'More Experiences? Yes or No '
    )
    if more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter Company: ')
        start_date = input('Start Date: ')
        to_date = input('To Date: ')
        experience = input('Describe your experience at '+ company + ': ')

        p.add_run(company + ' ' + '\n').bold = True
        p.add_run(start_date + '-' + to_date + '\n').italics =  True
        p.add_run(experience)
    else:
        break

# skills
document.add_heading('Skills')
skills = input('Enter Skill: ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'


while True:
    more_skills = input('More_skills? Yes or No ')
    if more_skills.lower() == 'yes':
        skills = input('Enter Skill: ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Temmys code'

document.save('cv.docx')